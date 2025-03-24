from datetime import date,datetime
from email import message
import openai
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
import comtypes
from reportlab.lib.styles import getSampleStyleSheet 
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Image,Spacer # type: ignore
from django.shortcuts import get_object_or_404, render, redirect
import os
import json
from django.http import JsonResponse
from django.contrib.auth import logout
from django.core.files.storage import FileSystemStorage
from django.db.models import Q
# from weasyprint import HTML
from docx import Document
from django.conf import settings
from .models import * #noqa
from django.contrib import messages
from django.core.mail import send_mail
from itertools import groupby
from django.core.mail.message import EmailMessage
import tempfile
from django.utils.timezone import now
import re

openai.api_key = settings.OPENAI_API_KEY


def convert_docx_to_json(docx_path, filename):
    """Convert a .docx file to JSON format and save it to media folder"""
    document = Document(docx_path)
    data = {"paragraphs": []}

    for para in document.paragraphs:
        data["paragraphs"].append(para.text)
    
    json_path = os.path.join(settings.MEDIA_ROOT, "json", filename + ".json")
    
    try:
        with open(json_path, "w") as json_file:
            json.dump(data, json_file, indent=4)
    except Exception as e:
        print(f"Error saving JSON file: {e}")
    
    return json.dumps(data, indent=4)
# Create your views here.


def jobseeker_login(request):
    if request.method == "POST":
        Email = request.POST.get("email")
        Password = request.POST.get("password")
        try:
            jobseeker = Jobseeker_Registration.objects.get(email=Email, password=Password)
            request.session['jobseeker_name'] = jobseeker.name
            request.session['jobseeker_id'] = jobseeker.id
            request.session['jobseeker_email'] = jobseeker.email
            request.session['jobseeker_phone'] = jobseeker.phone
            messages.success(request, "Login Successfully")
            return redirect('jobseeker_dashboard')  # Redirect to a dashboard or home page after login
        except Jobseeker_Registration.DoesNotExist:
            messages.error(request, "Invalid Email or Password")
    return render(request, 'jobseeker_login.html')

def extract_resume_details(content):
    """Extracts key details like name, skills, address, highest_qualification, job_preference, university name, date of birth, email, and phone number from the resume content."""
    details = {
        "name": "",
        "skills": "",
        "address": "",
        "highest_qualification": "",
        "passout_year": "",
        "percentage": "",
        "job_preference": "",
        "university": "",
        "dob": "",
        "email": "",
        "phone": ""
    }
    
    # Define patterns and keywords
    email_pattern = r"[a-zA-Z0-9+_.-]+@[a-zA-Z0-9.-]+"
    phone_pattern = r"\+?\d{10,15}"
    # passout_year_pattern = r"\b(19\d{2}|20\d{2})\b"
    percentage_pattern = r"(\d{1,2}\.\d{1,2}|\d{1,2})%"
    dob_patterns = [
        r"\b(\d{1,2})[-/ ](\d{1,2})[-/ ](\d{2,4})\b",  # Formats like DD/MM/YYYY, DD-MM-YYYY
        r"\b(\d{1,2})[-/ ](Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[-/ ](\d{2,4})\b"  # Formats with month names
    ]
    qualification_keywords = ["Bachelor's Degree","Integrated MCA","MCA", "Master of Computer Application", "PhD", "B.Sc", "M.Sc", "B.Tech Computer Science", "M.Tech Computer Science", "MBA"]
    job_preferences_keywords = ["Software Engineer", "Data Scientist", "Backend Developer", "Frontend Developer", "Project Manager"]
    university_keywords = ["University", "Institute", "College"]
    skills_keywords = ["Python", "Java", "C++", "Django", "SQL", "Machine Learning", "Artificial Intelligence", "React", "JavaScript", "HTML", "CSS", "Git", "Data Analytics"]
    address_keywords = ["State", "Country", "District"]
    indian_states = ["Andhra Pradesh", "Arunachal Pradesh", "Assam", "Bihar", "Chhattisgarh", "Goa", "Gujarat", "Haryana", "Himachal Pradesh", "Jharkhand", "Karnataka", "Kerala", "Madhya Pradesh", "Maharashtra", "Manipur", "Meghalaya", "Mizoram", "Nagaland", "Odisha", "Punjab", "Rajasthan", "Sikkim", "Tamil Nadu", "Telangana", "Tripura", "Uttar Pradesh", "Uttarakhand", "West Bengal"]
    
    # Use a set for skills to prevent duplicates
    skills_found = set()
    
    # Process content line-by-line
    lines = content.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Extract name: assume first occurrence of two or more capitalized words is the candidate's name
        if not details["name"]:
            name_match = re.match(r"^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)", line)
            if name_match:
                details["name"] = name_match.group(1)
        
        # Extract email
        if not details["email"]:
            email_match = re.search(email_pattern, line)
            if email_match:
                details["email"] = email_match.group()
        
        # Extract phone
        if not details["phone"]:
            phone_match = re.search(phone_pattern, line)
            if phone_match:
                details["phone"] = phone_match.group()
        
        # # Extract date of birth
        # if not details["dob"]:
        #     dob_match = re.search(dob_patterns, line, flags=re.IGNORECASE)
        #     if dob_match:
        #         details["dob"] = dob_match.group(1)
        # Extract date of birth and format to YYYY-MM-DD (for HTML date input)
        if not details["dob"]:
            for pattern in dob_patterns:
                dob_match = re.search(pattern, line, flags=re.IGNORECASE)
                if dob_match:
                    try:
                        if len(dob_match.groups()) == 3:
                            day, month, year = dob_match.groups()
                            if month.isdigit():
                                formatted_date = f"{int(year):04d}-{int(month):02d}-{int(day):02d}"
                            else:
                                month_num = {
                                    "Jan": "01", "Feb": "02", "Mar": "03", "Apr": "04", "May": "05", "Jun": "06", 
                                    "Jul": "07", "Aug": "08", "Sep": "09", "Oct": "10", "Nov": "11", "Dec": "12"
                                }.get(month[:3].capitalize(), "")
                                if month_num:
                                    formatted_date = f"{int(year):04d}-{month_num}-{int(day):02d}"
                                else:
                                    continue
                            details["dob"] = formatted_date
                            break
                    except ValueError:
                        continue
        
        # Extract address if line contains address-related keywords or Indian states
        if not details["address"]:
            if any(kw.lower() in line.lower() for kw in address_keywords) or any(state.lower() in line.lower() for state in indian_states):
                details["address"] = line
        
        # Extract highest qualification
        if not details["highest_qualification"]:
            for qualification in qualification_keywords:
                if qualification.lower() in line.lower():
                    details["highest_qualification"] = qualification
                    break
        
        # Extract percentage
        if not details["percentage"]:
            percentage_match = re.search(percentage_pattern, line)
            if percentage_match:
                details["percentage"] = percentage_match.group()
                
        # Extract job preference
        if not details["job_preference"]:
            for job in job_preferences_keywords:
                if job.lower() in line.lower():
                    details["job_preference"] = job
                    break
        
        # Extract university: store the entire line if it contains any keyword
        if not details["university"]:
            for uni_kw in university_keywords:
                if uni_kw.lower() in line.lower():
                    details["university"] = line
                    break
        
        # Accumulate skills from the line
        for skill in skills_keywords:
            if skill.lower() in line.lower():
                skills_found.add(skill)
    
    details["skills"] = ", ".join(sorted(skills_found))
    return details


def jobseeker_home(request):
    content = ""
    alert_message = ""
    fname = request.session.get('jobseeker_name')
    uid = request.session.get('jobseeker_id')
    jobseek_email = request.session.get('jobseeker_email')
    phone = request.session.get('jobseeker_phone')
    extracted_details = None
    show_modal = False

    if request.method == "POST":
        uploaded_file = request.FILES.get("word_file")

        # Check if a resume already exists for the user
        if jobseeker_resume.objects.filter(user_id=uid).exists():
            messages.error(request, "You have already uploaded a resume. You cannot upload another one.")
            return redirect("home")

        if uploaded_file:
            allowed_extensions = ('.doc', '.docx')
            file_extension = os.path.splitext(uploaded_file.name)[1].lower()
            if file_extension not in allowed_extensions:
                messages.error(request, "Invalid file format. Only .doc, .docx files are allowed.")
                return redirect("home")

            try:
                try:
                    jobseeker_profile.objects.get(user_id=uid)
                except jobseeker_profile.DoesNotExist:
                    show_modal = True

                documents_dir = os.path.join(settings.MEDIA_ROOT, "documents")
                os.makedirs(documents_dir, exist_ok=True)

                fs = FileSystemStorage(location=documents_dir)
                saved_filename = fs.save(uploaded_file.name, uploaded_file)
                file_path = fs.path(saved_filename)

                # Enhanced extraction based on file type
                if file_extension in ('.doc', '.docx'):
                    document = Document(file_path)
                    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
                    content = "\n".join(paragraphs)

                # Remove the file after processing
                os.remove(file_path)

                # Extract resume details using the improved extraction function
                extracted_details = extract_resume_details(content)
                essential_fields = ["name", "email", "phone", "address", "skills", "university"]
                if any(not extracted_details.get(field) for field in essential_fields):
                    messages.error(request, "Failed to extract essential details from your resume. Check your file and try again.")
                    return redirect("home")

                # Save the resume file in the database after successful extraction
                jobseeker_resumeobj = jobseeker_resume(file=uploaded_file, user_id=uid)
                jobseeker_resumeobj.save()

            except Exception as e:
                messages.error(request, f"Error processing document: {str(e)}")
                return redirect("home")
    resume = jobseeker_resume.objects.filter(user_id=uid)
    job=job_title.objects.all()
    jobseeker=""
    if(jobseeker_profile.objects.filter(user_id=uid).exists()):
        jobseeker=jobseeker_profile.objects.get(user_id=uid).profile_img
    context = {
        "resume_details": extracted_details,
        "alert_message": alert_message,
        "fname": fname,
        "email": jobseek_email,
        "phone": phone,
        "show_modal": show_modal,
        "resume": resume,
        "job":job,
        "jsdt":jobseeker
    }

    if not jobseeker_profile.objects.filter(name=fname).exists() and not jobseeker_resume.objects.filter(user=uid).exists():
        redirect('home')
    return render(request, "jobseeker_home.html", context)


def jobseeker_profile_update(request):
    userid = request.session.get('jobseeker_id')
    dob = ""
    phone = ""
    address = ""
    qualification = ""
    skills = ""
    job_preference = ""
    if request.method == "POST": 
        name = request.POST.get('name')
        email = request.POST.get('email')
        dob = request.POST.get('dob')
        phone = request.POST.get('phone')
        address = request.POST.get('Address')
        qualification = request.POST.get('Education')
        university = request.POST.get('University')
        percentage=request.POST.get('Percentage')
        passout=request.POST.get('passout')
        skills = request.POST.get('skills')
        image = request.FILES.get('image')
        job_preference = request.POST.get('job_preferences')
        # Fetch the first resume for the user (if multiple exist)
        resume = jobseeker_resume.objects.filter(user=userid).first()
        cv = resume.file if resume else None  # Handle case where no resume exists
        job_preference=job_title.objects.get(id=job_preference).job_title
        if jobseeker_profile.objects.filter(name=name).exists() and jobseeker_profile.objects.filter(email=email).exists():
            messages.error(request, "A Jobseeker with this name and/or email already exists. Please use different details.")
        else:
            jobseeker_obj = jobseeker_profile()
            jobseeker_obj.name = name
            jobseeker_obj.email = email
            jobseeker_obj.dob = dob
            jobseeker_obj.highest_qualification = qualification
            jobseeker_obj.percentage=percentage
            jobseeker_obj.passoutyear=passout
            jobseeker_obj.skills = skills
            jobseeker_obj.job_preference = job_preference
            jobseeker_obj.university = university
            jobseeker_obj.address = address
            jobseeker_obj.phone = phone
            jobseeker_obj.profile_img = image
            jobseeker_obj.user_id = userid
            jobseeker_obj.resume = cv
            jobseeker_obj.save()
            request.session['address'] = jobseeker_obj.address
            request.session['dob'] = jobseeker_obj.dob
            messages.success(request, "Updated profile successfully.Now you can use whole features of this portal.")
    return redirect('home')


def main(request):
    return render(request, 'index.html')

def contact_view(request):
    if request.method == "POST":
        name = request.POST.get("name")
        email = request.POST.get("email")
        message = request.POST.get("message")

        if name and email and message:
            # Send email
            send_mail(
                subject=f"New Contact Form Submission from {name}",
                message=f"From \nName: {name}\nEmail: {email}\n\n\n{message}",
                from_email=email,
                recipient_list=["jobai.prksolutions@gmail.com"],  # Replace with your email
                fail_silently=False,
            )

            messages.success(request, "Your message has been sent successfully!")
            return redirect("index")  # Redirect to the contact page after submission
        else:
            messages.error(request, "All fields are required.")

    return render(request, "index.html")  # Replace with your actual template file

def base(request):
    jbid = request.session.get('jobseeker_id')
    if jobseeker_profile.objects.filter(user=jbid).exists():
        jsdt = jobseeker_profile.objects.get(user=jbid)
        return render(request, 'base.html', jsdt)
    else:
        jsdt = None
    jobseeker = jobseeker_profile.objects.get(user=jbid)
    unread_notifications = JobNotification.objects.filter(jobseeker_profile=jobseeker, is_read=False).count()
    context={
        "jsdt":jsdt,"count":unread_notifications
    }
    return render(request, 'base.html', context)


def Register(request):
    name = ""
    email = ""
    password = ""
    if request.method == "POST": 
        name = request.POST.get('name')
        email = request.POST.get('email')
        password = request.POST.get('password')
        phone = request.POST.get('Phone')
        if Jobseeker_Registration.objects.filter(name=name).exists() or Jobseeker_Registration.objects.filter(email=email).exists():
            messages.error(request, "A Jobseeker with this name and/or email already exists.You can Login directly.")
        else:
            jobseeker_obj = Jobseeker_Registration()
            jobseeker_obj.name = name
            jobseeker_obj.email = email
            jobseeker_obj.password = password
            jobseeker_obj.phone = phone
            jobseeker_obj.save()
            messages.success(request, "Jobseeker registered successfully!.Email will be your Username")
            return render(request, 'jobseeker_register.html')
    return render(request, 'jobseeker_register.html')


def reset_password(request, user_id):
    try:
        user = Jobseeker_Registration.objects.get(pk=user_id)
    except Jobseeker_Registration.DoesNotExist:
        messages.error(request, "Invalid password reset link.")
        return redirect("forgot_password")

    if request.method == "POST":
        new_password = request.POST.get("get_password")
        confirm_password = request.POST.get("confirm_password")

        if new_password != confirm_password:
            messages.error(request, "Passwords do not match.")
        else:
            user.password = new_password 
           
            user.save()
            messages.success(request, "Your password has been reset successfully.")
            return redirect("jobseeker_login")  # Redirect after reset

    return render(request, "j_reset_password.html", {"user_id":user_id})


def Forgot_pwd(request):
    if request.method == "POST":
        email = request.POST.get("email")
        try:
            user = Jobseeker_Registration.objects.get(email=email)
            uid = user.id
            # Send email without link
            reset_url = f"{settings.SITE_URL}/reset_password/{uid}/"
            
            # Send the email
            send_mail(
                subject="Password Reset Request",
                message="",
                html_message=f"""Dear { user.name },\nClick the button below to reset your password:\n\n<br><a href='{reset_url}'><button style='background-color:blue;color:white;border:1px solid blue;font-weight:bold;border-radius:10px'>Reset Password</button></a>""",
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=[email],
                fail_silently=False,
            )
            messages.success(request, "Password reset instructions have been sent to your email.")
        except Jobseeker_Registration.DoesNotExist:
            messages.error(request, "Email not found. Please enter a registered email.")

    return render(request, "forgot-password.html")


def user_base(request):
    return render(request, 'base.html')


def jobseeker_logout(request):
    logout(request)
    request.session.clear()
    messages.success(request, "Logout Successfully")
    return render(request, 'jobseeker_login.html')


def company_logout(request):
    logout(request)
    request.session.clear()
    messages.success(request, "Logout Successfully")
    return render(request, 'company_login.html')


def search_job(request):
    # Get the jobseeker name from session and check if the profile exists
    fname = request.session.get('jobseeker_name')
    if not jobseeker_profile.objects.filter(name=fname).exists():
        return redirect('home')
    
    jobseeker = jobseeker_profile.objects.get(name=fname)
    jobs = company_joblist.objects.filter(Lastdate__gt=date.today())

    # Get search parameters from GET request
    search_query = request.GET.get('search', '')
    location_query = request.GET.get('location', '')
    
    # Filter by search query. For the foreign key job_title, traverse to its text field.
    if search_query:
        jobs = jobs.filter(
        Q(job_title__job_title__istartswith=search_query) | 
        Q(company__name__icontains=search_query) | 
        Q(job_type__icontains=search_query)
        )
    
    # Filter by location query
    if location_query:
        jobs = jobs.filter(location__icontains=location_query)
    unread_notifications = JobNotification.objects.filter(jobseeker_profile=jobseeker, is_read=False).count()
    context = {
        "fname": fname,
        "jobs": jobs,
        "jsdt": jobseeker,
        'count':unread_notifications
    }
    return render(request, 'search-job.html', context)

def main_search(request):
    jobs = company_joblist.objects.filter(Lastdate__gt=date.today())

    # Get search parameters from GET request
    search_query = request.GET.get('search', '')
    location_query = request.GET.get('location', '')
    
    # Filter by search query. For the foreign key job_title, traverse to its text field.
    if search_query:
        jobs = jobs.filter(
        Q(job_title__job_title__istartswith=search_query) | 
        Q(company__name__icontains=search_query) | 
        Q(job_type__icontains=search_query)
        )
    
    # Filter by location query
    if location_query:
        jobs = jobs.filter(location__icontains=location_query)
    context = {
        "jobs": jobs,
    }
    return render(request,'main-search.html',context)
def coverletter(request):
    job = job_title.objects.all()
    company = Company.objects.all()
    fname = request.session.get('jobseeker_name')
    email = request.session.get('jobseeker_email')
    phoneno = request.session.get('jobseeker_phone')
    if not jobseeker_profile.objects.filter(name=fname).exists():
        return redirect('home')

    if request.method == "POST":
        # Get form data
        name = request.POST.get("name", "Your Name")
        position = request.POST.get("job_title", "Job Title")
        company_id = request.POST.get("Companyname", "Company Name")
        hr_name = request.POST.get("hrname", "Hiring Manager")
        skills = request.POST.get("skills", "Your Skills")
        job_type = request.POST.get("jobtype", "Job Type")
        phone = request.POST.get("phone", "Phone Number")
        address = request.POST.get("Address")
        eligibility = request.POST.get('eligibility')
        ldate = date.today().strftime('%d-%m-%Y')
        try:
            company_obj = Company.objects.get(id=company_id)
            company_name = company_obj.name
            company_loc = company_obj.address
            job_position = job_title.objects.get(id=position)
            jobs = job_position.job_title
            job_id=jobs
        except Company.DoesNotExist:
            company_name = "Unknown Company"

        # OpenAI API Call for Cover Letter Generation
        prompt = f"""
Write a professional and engaging cover letter for a fresher job seeker named {name} , whose address is {address} ,email is {email},having qualification of '{eligibility}' and phone number is {phone}. 
The job seeker is applying for a {job_type} position as {jobs} at {company_name}, located at {company_loc}. 

Address the letter to {hr_name}, incorporating the company's full address. Highlight relevant skills such as {skills} and ensure the letter is formatted in a common professional style without placeholders like [Your Name] or [Your Address]. 

Include the date ({ldate}) at the beginning and conclude with a proper closing, using the candidate’s full name instead of placeholders. Keep the letter concise, engaging, and directly relevant to the job opportunity.
"""
        
        response = openai.ChatCompletion.create(# pylint: disable=undefined-variable

             model="gpt-4o-mini",
  store=True,
  messages=[
    {"role": "user", "content": prompt}
  ]
        )
        request.session['company_id']=company_id
        request.session['highest_qualification']=eligibility
        cover_letter = response["choices"][0]["message"]["content"].strip()
        jsdt = jobseeker_profile.objects.get(email=email)
        unread_notifications = JobNotification.objects.filter(jobseeker_profile=jsdt, is_read=False).count()
        return render(request, "cover-letter.html", {
            "fname": fname,
            "email": email,
            "cover_letter": cover_letter,
            "company": company,
            "phone":phoneno,
            "jsdt":jsdt,
            "count":unread_notifications,
            "job":job,
            "job_id":job_id
        })
    jsdt = jobseeker_profile.objects.get(email=email)
    unread_notifications = JobNotification.objects.filter(jobseeker_profile=jsdt, is_read=False).count()
    return render(request, "cover-letter.html", {
        "fname": fname,
        "company": company,
        "email": email,
        "phone":phoneno,
        "job":job,
        "jsdt":jsdt,
        "count":unread_notifications
    })

def send_cover_letter_email(request):
    fname=request.session.get('jobseeker_name')
    # highest_qualification=request.session.get('highest_qualification')
    
    if request.method == "POST":
        cid = request.session.get("company_id")
        pdf_file = request.FILES.get("pdf")
        # job_title_id=request.POST.get('job_id')
        # email=Company.objects.get(id=cid).email
        email=request.session.get('jobseeker_email')
        if not cid or not pdf_file:
            return JsonResponse({"message": "Missing email or PDF!"}, status=400)
        # job=job_title.objects.get(id=job_title_id).job_title
        # Save the uploaded PDF to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            for chunk in pdf_file.chunks():
                temp_pdf.write(chunk)

            temp_pdf_path = temp_pdf.name

        # Send email
        email_message = EmailMessage(
            subject=f"{ fname }'s Cover Letter Generated in JobAi",
            body=f"""
Dear {fname},

I hope this email finds you well.

We are pleased to inform you that your cover letter has been successfully generated by the JobAi portal. Please find the attached cover letter, tailored specifically to the information you've provided. We hope this will help you in your job search and support you in landing your next opportunity.

Best of luck with your applications, and don't hesitate to reach out if you need any further assistance.

Best regards,
The JobAi Team
""",
            from_email=settings.DEFAULT_FROM_EMAIL,  
            to=[email],
        )
        email_message.attach_file(temp_pdf_path)
        email_message.send()
        return JsonResponse({"message": "Cover Letter sent successfully!"})
    return JsonResponse({"message": "Invalid request method"}, status=400)
def settings_view(request): 
    jbid = request.session.get('jobseeker_id')
    try: 
        jsdt = jobseeker_profile.objects.get(user_id=jbid)
        fname = request.session.get('jobseeker_name')
        jobseek_email = request.session.get('jobseeker_email') 
        phone = request.session.get('jobseeker_phone')
        unread_notifications = JobNotification.objects.filter(jobseeker_profile=jsdt, is_read=False).count()
        # address=request.session.get("address")
        context = {
            'jsdt':jsdt,
                "fname":fname,
                "email":jobseek_email,
                "phone":phone,
                "count":unread_notifications
            }
        return render(request, 'settings_view.html', context)
    except:
        
        jsdt = Jobseeker_Registration.objects.filter(id=jbid)
        fname = request.session.get('jobseeker_name')
        if not jobseeker_profile.objects.filter(name=fname).exists():
            return redirect('home')
        context = {
            'jsdt':jsdt,
                "fname":fname,
              
            }
        return render(request, 'settings_view.html', context)


def support(request):
    fname = request.session.get('jobseeker_name')
    if not jobseeker_profile.objects.filter(name=fname).exists():
            return redirect('home')
    jobseeker = jobseeker_profile.objects.get(name=fname)
    unread_notifications = JobNotification.objects.filter(jobseeker_profile=jobseeker, is_read=False).count()
    return render(request, 'support.html', {"fname":fname, "jsdt":jobseeker,"count":unread_notifications})


def mockinterview(request):
    fname = request.session.get('jobseeker_name')
    if not jobseeker_profile.objects.filter(name=fname).exists():
            return redirect('home')
    jobs = job_title.objects.all()
    if request.method == "POST":
        job_id = request.POST.get("job_title")
        user_answer = request.POST.get("answer", None)

        if not job_id:
            return JsonResponse({"error": "Job title is required"}, status=400)

        try:
            job = job_title.objects.get(id=job_id)
        except job_title.DoesNotExist:
            return JsonResponse({"error": "Invalid job title"}, status=400)

        if user_answer is None:
            # Generate exactly 5 interview questions for the given job title
            prompt = f"Provide exactly 5 professional interview questions for a {job.job_title} role. Do not include any introduction or explanation. List them one after another."
            response = openai.ChatCompletion.create(# pylint: disable=undefined-variable

                model="gpt-4o-mini",
                messages=[{"role": "system", "content": prompt}]
            )

            content = response["choices"][0]["message"]["content"].strip()
            # Split questions by newline and filter out empty strings
            questions = [q for q in content.split("\n") if q.strip()]
            # Ensure that exactly 5 questions are returned
            if len(questions) != 5:
                return JsonResponse({"error": "Failed to generate exactly 5 fresher level interview questions. Please try again."}, status=400)
            return JsonResponse({"questions": questions})

        else:
            # Evaluate the user answer
            prompt = f"Evaluate this job interview response for a {job.job_title} role and provide constructive feedback,score,sample answer for the questions for beginners. Do not include any introduction or pleasantries. Answer concisely.\n\nAnswer: {user_answer}"
            response = openai.ChatCompletion.create(# pylint: disable=undefined-variable
                model="gpt-4o-mini",
                messages=[{"role": "system", "content": prompt}]
            )
            feedback = response["choices"][0]["message"]["content"].strip()
            return JsonResponse({"feedback": feedback})
    jobseeker = jobseeker_profile.objects.get(name=fname)
    unread_notifications = JobNotification.objects.filter(jobseeker_profile=jobseeker, is_read=False).count()
    return render(request, 'mock-interview.html', {'job': jobs, "fname":fname, "jsdt":jobseeker,"count":unread_notifications})

def parse_list(text):
    """
    Parse a comma-separated string into a set of lowercase trimmed items.
    """
    if text:
        return set([item.strip().lower() for item in text.split(",") if item.strip()])
    return set()

def match_jobs(jobseeker):
    # Function to extract numeric percentage value from string
    def extract_percentage(value):
        try:
            return float(value.strip('%')) if value else 0.0
        except ValueError:
            return 0.0
    # Parse jobseeker skills and qualification
    seeker_skills = parse_list(jobseeker.skills)
    seeker_qual = jobseeker.highest_qualification.strip().lower() if jobseeker.highest_qualification else ""
    seeker_percentage = extract_percentage(jobseeker.percentage) if jobseeker.percentage else 0.0
    # Fetch all jobs (this could be optimized with a filter, but for now we'll do it in Python)
    all_jobs = company_joblist.objects.all()
    matching_jobs = []
    for job in all_jobs:
        job_skills = parse_list(job.skills_required)
        job_percentage_criteria = extract_percentage(job.percent_criteria)
        
        # Check for any common skills
        if not seeker_skills.intersection(job_skills):
            continue
        
        # Check eligibility: if job.highest_qualification exists, see if it matches the seeker qualification
        if job.highest_qualification:
            # Split eligibility criteria if multiple (assuming comma-separated)
            eligibility_set = parse_list(job.highest_qualification)
            if seeker_qual and seeker_qual not in eligibility_set:
                continue
            # Check if jobseeker's percentage meets or exceeds the job's percentage criteria
        if seeker_percentage < job_percentage_criteria:
            continue
        matching_jobs.append(job)
    return matching_jobs

def auto_apply_jobs(jobseeker, max_apply=4):
    today = now().date()
    
    # Get job seeker's preference (convert to lowercase for case-insensitive matching)
    job_preference = jobseeker.job_preference.lower() if jobseeker.job_preference else ""

    # Get applications already made today
    today_applied_jobs = JobApplication.objects.filter(jobseeker=jobseeker, applied_at__date=today)
    today_applied_jobs_count = today_applied_jobs.count()

    # **Enforce max 4 applications upfront**
    if today_applied_jobs_count >= max_apply:
        print(f"⚠️ Jobseeker {jobseeker.id} already applied for {today_applied_jobs_count} jobs today. No more applications allowed.")
        return 0, []

    # Get list of job IDs already applied for today (prevents duplicate applications for the same job)
    applied_job_ids = set(today_applied_jobs.values_list('company_joblist_id', flat=True))

    # Get matching jobs based on skills and eligibility
    matching_jobs = match_jobs(jobseeker)

    # Exclude jobs already applied for today and ensure job is still open
    available_jobs = [
        job for job in matching_jobs 
        if job.id not in applied_job_ids  # Prevent duplicate job applications
        and job.Lastdate >= today  # Ensure job is still open
    ]

    # **Sort jobs based on priority**
    def job_sort_key(job):
        is_preferred = job_preference in job.job_title.job_title.lower()
        return (
            not is_preferred,  # Prioritize job preference
            job.Lastdate,       # Prioritize nearest deadline
            job.dateofpublish or ""  # Prioritize recently published jobs
        )

    sorted_jobs = sorted(available_jobs, key=job_sort_key)

    applications_made = 0
    applied_jobs_list = []

    for job in sorted_jobs:
        # **Strictly enforce max-apply limit**
        if (today_applied_jobs_count + applications_made) >= max_apply:
            print(f"Max applications ({max_apply}) reached. Stopping further applications.")
            break  

        # **Final check before applying**
        if JobApplication.objects.filter(jobseeker=jobseeker, company_joblist=job).exists():
            print(f"Already applied to {job.job_title.job_title} at {job.company.name}. Skipping...")
            continue

        try:
            # Apply for the job
            JobApplication.objects.create(jobseeker=jobseeker, company_joblist=job)
            applications_made += 1  # Track successful applications
            applied_jobs_list.append(job)

            # Add this job to the exclusion list for future checks
            applied_job_ids.add(job.id)
            print(f"Applied to: {job.job_title.job_title} at {job.company.name}")

        except Exception as e:
            print(f"Error applying for {job.job_title.job_title}: {e}")
            continue

    # Send email notification if applications were made
    job_details_html = "".join([
    "<tr>"
    f"<td>🔹 <b>{job.job_title.job_title}</b></td>"
    f"<td>🏢 {job.company.name}</td>"
    f"<td>📍 {job.location}</td>"
    f"<td>🏷️ {job.job_type}</td>"
    f"<td>📅 {today}</td>"
    f"<td>⏳ {job.Lastdate}</td>"
   "</tr>"
    for job in applied_jobs_list
])

    email_body_html = f"""
<p>Dear {jobseeker.name},</p>
<p>You have successfully applied to the following jobs today via <b>JobAi</b>:</p>
<table border='1' cellspacing='0' cellpadding='5' style='border-collapse: collapse; text-align: left; width: 100%;'>
    <tr style='background-color: #f2f2f2;'>
        <th>Job Title</th>
        <th>Company</th>
        <th>Location</th>
        <th>Job Type</th>
        <th>Application Date</th>
        <th>Deadline</th>
    </tr>
    {job_details_html}
</table>
<p>🚀 <b>Next Steps:</b></p>
<ul>
    <li>📌 Check your application status in your <a href='http://127.0.0.1:8000/jobseeker_login/'>JobAi Profile</a>.</li>
    <li>📌 Prepare for interviews with our <b>AI Mock Interview Tool</b>.</li>
</ul>
<p>Thank you for using <b>JobAi</b>. We wish you success!</p>
<p>Best Regards,<br><b>JobAi Team</b><br>📧 support@jobai.com | 🌐 <a href='http://127.0.0.1:8000/jobseeker_login/'>www.jobai.com</a></p>
"""
    if applications_made!=0:
        send_mail(
    subject="Job Application Confirmation – Your Applications for Today",
    message="This is an HTML email. Please enable HTML to view the content properly.",
    from_email="jobai.prksolutions@gmail.com",
    recipient_list=[jobseeker.email],
    fail_silently=False,
    html_message=email_body_html
)
    return applications_made, applied_jobs_list


def autoapply(request):
    # Use session to retrieve jobseeker email instead of just name
    fname=request.session.get('jobseeker_name')
    jobseeker_email = request.session.get('jobseeker_email')
    if not jobseeker_email or not jobseeker_profile.objects.filter(email=jobseeker_email).exists():
        messages.error(request, "Jobseeker profile not found.")
        return redirect('home')

    jobseeker = jobseeker_profile.objects.get(email=jobseeker_email)
    recommended_jobs = match_jobs(jobseeker)
    recommended_jobs = [job for job in recommended_jobs if job.Lastdate >= now().date()]
    recommended_jobs=sorted(recommended_jobs, key=lambda x: x.Lastdate)
    # On POST request, perform auto-apply action
    if request.method == "POST":
        applications, _ = auto_apply_jobs(jobseeker, max_apply=4)
        if applications:
            messages.success(request, f"Auto-applied to {applications} job(s) successfully!")
        else:
            messages.info(request, "No new matching jobs available for auto-apply.")
        redirect('auto-apply')
    
    unread_notifications = JobNotification.objects.filter(jobseeker_profile=jobseeker, is_read=False).count()

    context = {
        "jobseeker_email": jobseeker_email,
        "jsdt": jobseeker,
        "fname":fname,
        "recommended_jobs": recommended_jobs,
        "count":unread_notifications
    }
    return render(request, "auto-apply.html", context)

def applied_jobs(request):
    fname = request.session.get('jobseeker_name')
    if not fname or not jobseeker_profile.objects.filter(name=fname).exists():
        return redirect('home')
    profile=jobseeker_profile.objects.get(name=fname)
    profile_id=profile.id
    jobseeker = jobseeker_profile.objects.get(name=fname)
    jobs=JobApplication.objects.filter(jobseeker_id=profile_id)
    jobs=sorted(jobs, key=lambda x: x.id)
    unread_notifications = JobNotification.objects.filter(jobseeker_profile=jobseeker, is_read=False).count()
    context = {
        "fname": fname,
        "jsdt": jobseeker,
        "jobs":jobs,
        "count": unread_notifications,
    }
    return render(request,'applied-jobs.html',context)
def shortlisted(request):
    fname = request.session.get('jobseeker_name')
    if not fname or not jobseeker_profile.objects.filter(name=fname).exists():
        return redirect('home')
    profile=jobseeker_profile.objects.get(name=fname)
    profile_id=profile.id
    jobseeker = jobseeker_profile.objects.get(name=fname)
    jobs=JobApplication.objects.filter(jobseeker_id=profile_id,status="Accepted")
    jobs=sorted(jobs, key=lambda x: x.id)
    unread_notifications = JobNotification.objects.filter(jobseeker_profile=jobseeker, is_read=False).count()
    context = {
        "fname": fname,
        "jsdt": jobseeker,
        "jobs":jobs,
        "count": unread_notifications,
    }
    return render(request,'jobseeker_shortlisted.html',context)

def user_type(request):
    return render(request, 'user-type.html')


def company_login(request):
    if request.method == "POST":
        Email = request.POST.get("email")
        Password = request.POST.get("password")
        try: 
            company = Company.objects.get(email=Email, password=Password) 
            request.session['company_id'] = company.id
            request.session['company_name'] = company.name
            request.session['company_email'] = company.email
            # request.session['company_address']=company.address
            # request.session['company_type']=company.company_type_id
            messages.success(request, "Login Successfully")
            return redirect('company_dashboard')  # Redirect to a dashboard or home page after login
        except Company.DoesNotExist:
            messages.error(request, "Invalid Email or Password")
    return render(request, 'company_login.html')


def company_registration(request):
    if request.method == "POST":
        name = request.POST.get('company_name')
        email = request.POST.get('email')
        password = request.POST.get('password')
        company_type = request.POST.get("CompanyType")  # This returns a string
        company_address = request.POST.get('Address')
        comp_logo = request.FILES.get('company_logo')
        # Check if a company with the same name or email already exists
        if Company.objects.filter(name=name).exists() or Company.objects.filter(email=email).exists():
            messages.error(request, "A company with this name or email already exists. You can Login.")
        else:
            # Create and save the company
            companyobj = Company()
            companyobj.name = name
            companyobj.password = password
            companyobj.email = email
            companyobj.company_type_id = company_type
            companyobj.address = company_address
            companyobj.profile_img = comp_logo
            companyobj.save()
            messages.success(request, "Company registered successfully!.Email will be your Username")
            return redirect('company_registration')
    company = Company_Type_Master.objects.all()
    context = {
        'company':company
    }
    return render(request, 'company_registration.html', context)

def change_password(request):
    fname = request.session.get('jobseeker_name')
    
    if not fname or not jobseeker_profile.objects.filter(name=fname).exists():
        return redirect('home')
    
    jobseeker = jobseeker_profile.objects.get(name=fname)
    unread_notifications = JobNotification.objects.filter(jobseeker_profile=jobseeker, is_read=False).count()
    user=Jobseeker_Registration.objects.get(name=fname)
    exist_pwd=user.password
    if request.method == "POST":
        old_password = request.POST.get('old_password')
        new_password = request.POST.get('get_password')
        confirm_password = request.POST.get('confirm_password')
        
        if exist_pwd != old_password:
            messages.error(request, "Old password is incorrect.")
        elif new_password != confirm_password:
            messages.error(request, "New passwords do not match.")
        elif len(new_password) < 4:
            messages.error(request, "Password must be at least 4 characters long.")
        elif exist_pwd==old_password and exist_pwd==new_password:
            messages.warning(request,"No changes in Password detected")
        else:
            userobj=Jobseeker_Registration.objects.get(name=fname)
            userobj.password=confirm_password
            userobj.save()
            messages.success(request, "Your password has been updated successfully!")

    context = {"fname": fname, "jsdt": jobseeker,"count":unread_notifications}
    return render(request, 'change_pwd.html', context)


def company_change_password(request):
    cname = request.session.get('company_name')
    cid = request.session.get('company_id')
    comp = Company.objects.get(id=cid)
    exist_pwd=comp.password
    if request.method == "POST":
        old_password = request.POST.get('old_password')
        new_password = request.POST.get('get_password')
        confirm_password = request.POST.get('confirm_password')
        
        if exist_pwd != old_password:
            messages.error(request, "Old password is incorrect.")
        elif new_password != confirm_password:
            messages.error(request, "New passwords do not match.")
        elif exist_pwd==old_password and exist_pwd==new_password:
            messages.warning(request,"No changes in Password detected")
        elif len(new_password) < 4:
            messages.error(request, "Password must be at least 4 characters long.")
        else:
            compobj=Company.objects.get(name=cname)
            compobj.password=confirm_password
            compobj.save()
            messages.success(request, "Your password has been updated successfully!")
    # 
    context = {
            # "jobs": expiredyet,
            "cname": cname,
            "compdt":comp
    }
    return render(request, 'company_change_pwd.html', context)

def delete_user_profile(request):
    fname = request.session.get('jobseeker_name')
    
    if not fname or not jobseeker_profile.objects.filter(name=fname).exists():
        return redirect('home')
    if(JobApplication.objects.filter(jobseeker__name=fname).exists()):
        JobApplication.objects.filter(jobseeker__name=fname).delete()
        
    # Get the jobseeker profile
    jobseeker = jobseeker_profile.objects.get(name=fname)

    # Delete profile image if exists
    if jobseeker.profile_img:
        profile_img_path = os.path.join(settings.MEDIA_ROOT, str(jobseeker.profile_img))
        if os.path.exists(profile_img_path):
            os.remove(profile_img_path)

    # Delete resume file if exists
    if jobseeker.resume:
        resume_path = os.path.join(settings.MEDIA_ROOT, str(jobseeker.resume))
        if os.path.exists(resume_path):
            os.remove(resume_path)

    # Delete the jobseeker profile
    jobseeker.delete()
    jobseeker_resume.objects.get(user__name=fname).delete()
    return redirect('settings_view')
    

def delete_job(request): 
    cname = request.session.get('company_name')
    job_id = request.GET.get("ids")
    messages.success(request, f"jobid:{str(job_id)}")
    comp = Company.objects.get(name=cname)
    if not job_id:
            messages.error(request, "Job ID is missing in request.")
            return render(request, 'company_jobs.html', {'cname': cname})
        
    try:
        # Delete all job applications related to this job
        job_applications = JobApplication.objects.filter(company_joblist__company__name=cname, company_joblist__id=job_id)
        if job_applications.exists():
            job_applications.delete()

        # Delete the job listing itself
        job = get_object_or_404(company_joblist, id=job_id)
        job.delete()
        messages.success(request, "Job deleted successfully")
        return redirect('job_listing')
    except Exception as e:
            messages.error(request, f"Error deleting job: {str(e)}")
    else: 
        messages.error(request, "Attempt for deleting job item failed")
        
    return render(request, 'company_jobs.html', {'cname': cname,"compdt":comp})

def convert_doc_to_docx(doc_path):
    """Converts a DOC file to DOCX using comtypes (Windows only)."""
    word = comtypes.client.CreateObject('Word.Application')#pylint:disable=undefined-variable
    doc = word.Documents.Open(doc_path)
    docx_path = doc_path + "x"
    doc.SaveAs(docx_path, FileFormat=16)  # 16 = docx format
    doc.Close()
    word.Quit()
    return docx_path

def extract_text_from_docx(docx_path):
    """Extracts text from a DOCX file."""
    doc = Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_resume_text(profile_id):
    """Fetches the resume path from the database and extracts text."""
    jobseeker = get_object_or_404(jobseeker_profile, id=profile_id)
    resume_path = os.path.join(settings.MEDIA_ROOT, str(jobseeker.resume))
    
    if resume_path.endswith(".docx"):
        return extract_text_from_docx(resume_path)
    elif resume_path.endswith(".doc"):
        docx_path = convert_doc_to_docx(resume_path)
        return extract_text_from_docx(docx_path)
    
    return None

def calculate_ats_score(resume_text, job_description):
    prompt = f"""
    You are an ATS resume screening AI. Evaluate the resume against the job description.

    **Scoring Criteria**:
    - **Skills and Qualifications match(50%)**
    - **Formatting & Readability (15%)**
    - **Grammar & Clarity (40%)**
    
    minimum 1/3 of appliication must be accepted so increase ats score of 3 for same job
    
    **Resume**:
    {resume_text}

    **Job Description**:
    {job_description}
    
    Provide:
    1. **ATS Score** (out of 100)
    2. **Improvement Suggestions** (if score < 50)
    """

    response = openai.ChatCompletion.create( #pylint: disable=undefined-variable
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0
    )
    
    ats_result = response["choices"][0]["message"]["content"]
    ats_score = int(re.search(r'\b(\d{1,3})\b', ats_result).group(1))
    return ats_score

def generate_admit_card(jobseeker, job_title, application_id):
    """Generates an Admit Card PDF with jobseeker details, company logo, and profile image."""
    
    # File Path for PDF
    file_name = f"{jobseeker.name.replace(' ', '_')}_admit_card.pdf"
    file_path = os.path.join(settings.MEDIA_ROOT, "admit_cards", file_name)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    application = get_object_or_404(JobApplication, id=application_id)

    # Create PDF Document
    pdf = SimpleDocTemplate(file_path, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()

    # Title
    title = Paragraph("<b>Aptitude Test Admit Card</b>", styles['Title'])
    
    # Load Company Logo
    logo_path = application.company_joblist.company.profile_img.path if application.company_joblist.company.profile_img else None
    company_logo = Image(logo_path, width=120, height=60) if logo_path and os.path.exists(logo_path) else Paragraph("<b>[Company Logo Not Available]</b>", styles['Normal'])
    
    # Load Jobseeker Profile Image
    profile_path = jobseeker.profile_img.path if jobseeker.profile_img else None
    applicant_image = Image(profile_path, width=100, height=100) if profile_path and os.path.exists(profile_path) else Paragraph("<b>[Applicant Image Not Available]</b>", styles['Normal'])
    
    # Table for Header with Company Logo & Title Alignment
    header_table = Table([[company_logo, title]], colWidths=[130, 400])
    header_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),  # Align logo left
        ('ALIGN', (1, 0), (1, 0), 'CENTER')  # Align title center
    ]))
    elements.append(header_table)
    elements.append(Spacer(1, 20))  # Space between title and next section
    
    # Jobseeker Image Alignment (Top Right)
    applicant_image_table = Table([[applicant_image]], colWidths=[100], rowHeights=[100])
    applicant_image_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),  # Align jobseeker image to the right
    ]))
    elements.append(applicant_image_table)
    elements.append(Spacer(1, 20))
    
    # Jobseeker Information Table
    data = [
        ["Applicant Name:", jobseeker.name],
        ["Email:", jobseeker.email],
        ["Phone:", jobseeker.phone],
        ["Qualification:", jobseeker.highest_qualification],
        ["Job Title:", job_title],
        ["Company:", application.company_joblist.company.name],
        ["Status:", "Selected"]
    ]

    table = Table(data, colWidths=[150, 300])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
    ]))
    elements.append(table)
    
    # Final Instructions
    instructions = Paragraph("""
        <br/><b>Instructions:</b><br/>
        Please bring this admit card along with required documents on the test day.
    """, styles['Normal'])
    elements.append(instructions)
    
    # Generate PDF
    pdf.build(elements)
    
    return file_path

def auto_process_application(request, application_id):
    """Automatically processes the job application based on ATS score."""
    try:
        application = get_object_or_404(JobApplication, id=application_id)
        jobseeker = get_object_or_404(jobseeker_profile, id=application.jobseeker_id)
        resume_text = extract_resume_text(application.jobseeker_id)
        if not resume_text:
            return {"error": "Failed to extract resume text."}
        else:
            print(resume_text)
        # Skip processing if already completed
        if application.ats_score is not None:
            return JsonResponse({
                "status": application.status,
                "ats_score": application.ats_score
            })

        #  Calculate ATS Score
        ats_score = calculate_ats_score(resume_text,application.company_joblist.job_description)


        #  Save ATS Score and Status in Database
        application.ats_score = ats_score
        application.save()

        # Determine Acceptance or Rejection
        if ats_score >= 55:
            application.status = "Accepted"
            status_text = "Accepted"
            status_class = "text-success fw-bold"
            
            # Generate Admit Card
            admit_card_path = generate_admit_card(jobseeker, application.company_joblist.job_title.job_title,application_id)
            
            # Email Content for Accepted Applicants
            email_subject = f"You're eligible to attend for Aptitude test for position {application.company_joblist.job_title.job_title} at {application.company_joblist.company.name}"
            email_body = f"""
Dear {jobseeker.name},

We are thrilled to inform you that your application for the position of {application.company_joblist.job_title.job_title} has been Accepted! 
 
Your admit card is attached with this email.

Next Steps:  
- Kindly review the admit card and take print out.
- Keep your admit card with you while come for aptitude test. 
- Ensure your participation in aptitude test.
 

We hope you prepare well for the test .The more details regarding Date,Time,Venue of exam will be informed shortly.

Best Regards,  
Recruitment Team
"""
            email = EmailMessage(email_subject, email_body, "noreply@company.com", [jobseeker.email])
            email.attach_file(admit_card_path)
            email.send()

        else:
            application.status = "Rejected"
            status_text = "Rejected"
            status_class = "text-danger fw-bold"

            # Email Content for Rejected Applicants
            email_subject = f"⚠️ Application Update - {application.company_joblist.job_title.job_title}"
            email_body = f"""
Dear {jobseeker.name},

We regret to inform you that your application for the position of {application.company_joblist.job_title.job_title} has been Rejected.  

We encourage you to improve your resume and apply for future opportunities with us.

Wishing you all the best in your job search!

Best Regards,  
Recruitment Team
"""
            email = EmailMessage(email_subject, email_body, "noreply@company.com", [jobseeker.email])
            email.send()

        # Save ATS Score & Status
        application.ats_score = ats_score
        application.save()
        if ats_score>100:
            messages.error(request,"Error occurred")
        # Return JSON Response for Frontend Update
        return JsonResponse({
            "status": application.status,
            "status_text": status_text,
            "status_class": status_class,
            "ats_score": ats_score
        })

    except JobApplication.DoesNotExist:
        return JsonResponse({"error": "Application not found"}, status=404)
def company_base(request):
    return render(request,'company_base.html')
def company_type(request):
    if request.method == "POST":
        company_type = request.POST.get('ctype')
        company_typeobj = Company_Type_Master()
        company_typeobj.company_type = company_type
        company_typeobj.save()
        messages.success(request, "Company type registered successfully!")
        return render(request, 'company_type.html')
    return render(request, 'company_type.html')


def jobs(request):
    if request.method == "POST":
        job_title1 = request.POST.get('job')
        company_typeobj = job_title()
        company_typeobj.job_title = job_title1
        company_typeobj.save()
        messages.success(request, "Job title added successfully!")
        return render(request, 'job_position.html')
    return render(request, 'job_position.html')


def company_reset_pwd(request, comp_id):
    try:
        comp = Company.objects.get(pk=comp_id)
    except Company.DoesNotExist:
        messages.error(request, "Invalid password reset link.")
        return redirect("forgot_password")

    if request.method == "POST":
        new_password = request.POST.get("get_password")
        confirm_password = request.POST.get("confirm_password")

        if new_password != confirm_password:
            messages.error(request, "Passwords do not match.")
        else:
            comp.password = new_password  # This should be hashed (use Django's set_password)
            comp.save()
            messages.success(request, "Your password has been reset successfully.")
            return redirect("company_login")  # Redirect after reset

    return render(request, "company_reset_password.html", {"comp_id":comp_id})


def company_forgot_password(request):
    if request.method == "POST":
        email = request.POST.get("email")
        try:
            cmpny = Company.objects.get(email=email)
            uid = cmpny.id
            # Setup link
            reset_url = f"{settings.SITE_URL}/company_reset_password/{uid}/"
            
            # Send the email
            send_mail(
                subject="Password Reset Request",
                message="",
                html_message=f"""\nClick the button below to reset your password:\n\n<br><a href='{reset_url}'><button style='background-color:blue;color:white;border:1px solid blue;font-weight:bold;border-radius:10px'>Reset Password</button></a>""",
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=[email],
                fail_silently=False,
            )
            
            messages.success(request, "Password reset instructions have been sent to your email.")
        except Company.DoesNotExist:
            messages.error(request, "Email not found. Please enter a registered email.")
    return render(request, 'company_forgot_pwd.html')


def company_dashboard(request):
    try:
        cname = request.session.get('company_name')
        cid = request.session.get('company_id')
        comp = Company.objects.get(id=cid)
        # Fetch all job listings for the logged-in company
        compjobdt = company_joblist.objects.filter(company_id=cid)
        participants=JobApplication.objects.filter(company_joblist__company_id=cid)
        latest_applications=participants.order_by('-id')[:3]
        pcount=participants.count()
        pendingapplications=JobApplication.objects.filter(company_joblist__company_id=cid,status="Applied").count()
        shortlisted=JobApplication.objects.filter(company_joblist__company_id=cid,status="Accepted").count()
        rejected=JobApplication.objects.filter(company_joblist__company_id=cid,status="Rejected").count()
        # If no jobs exist, show the default page
        if not compjobdt.exists():
            count = 0
            return render(request, 'company_dashboard.html', {"cname": cname, "count":count, "compdt":comp,"pcount":pcount,"participants":latest_applications}) # type: ignore
        
        # Get expired jobs
             
        count = compjobdt.count()  # Count jobs for the logged-in company
        context = {
            # "jobs": expiredyet,
            "cname": cname,
            "count":count,
            "compdt":comp,
            "pcount":pcount,
            "participants":latest_applications,
            "pending_applications":pendingapplications,
            "shortlisted_applications":shortlisted,
            "rejected_applications":rejected
        }
        return render(request, 'company_dashboard.html', context)
    except Exception as e:
        # Handle unexpected errors
        return render(request, 'company_dashboard.html', {"cname": cname, "compdt":comp, "error": str(e)})

def rejected(request):
    cid = request.session.get('company_id')
    cname = request.session.get('company_name')
    company = Company.objects.get(id=cid)
    Applicants=JobApplication.objects.filter(company_joblist__company_id=cid,status="Rejected")
    # Sort the list of applicants by job title
    Applicants = sorted(Applicants, key=lambda x: x.company_joblist.job_title.job_title)
    # Group applicants by job title
    grouped_applicants = {}
    for job_title, group in groupby(Applicants, key=lambda x: x.company_joblist.job_title.job_title):
            grouped_applicants[job_title] = list(group)
    
    context = {
        "cname":cname,
        "compdt":company,
        "Applicants":Applicants,
        "grouped_applicants":grouped_applicants,
       
    }
    return render(request,'rejected.html',context)

def company_settings(request):
    cid = request.session.get('company_id')
    compdt = Company.objects.get(id=cid)
    cmail = compdt.email
    cname = compdt.name
    cloc = compdt.address
    ctype = compdt.company_type.company_type
    company = Company.objects.get(id=cid)
    
    context = {
            "cname":cname, "cmail":cmail, "cloc":cloc,
            "ctype":ctype, "cid":cid, "compdt":company
        } 
    if request.method == "POST":
        comp_logo = request.FILES.get('company_logo')
        comp = Company.objects.get(id=cid)
        if comp_logo: 
            comp.profile_img = comp_logo
            comp.save()
            messages.success(request, "Logo Uploaded successfully!!")    
        return render(request, 'company_settings_view.html', context)
    return render(request, 'company_settings_view.html', context)


def company_jobs(request):
    try:
        cname = request.session.get('company_name')
        cid = request.session.get('company_id')
        company = Company.objects.get(id=cid)
        # Fetch all job listings for the logged-in company
        compjobdt = company_joblist.objects.filter(company=cid)
        # If no jobs exist, show the default page
        if not compjobdt.exists():
            return render(request, 'company_jobs.html', {"cname": cname, "compdt":company})
        
        # Iterate through each job for the company
        jobs_list = []
        for job in compjobdt:
            job_info = {
                "job_id": job.id,
                # "job_number": job.job_number,
                "job_title": job.job_title.job_title,
                "job_description": job.job_description,
                "eligibility":(job.highest_qualification),
                "skills": job.skills_required,
                "location": job.location,
                "job_type": job.job_type,
                "last_date": job.Lastdate,
                "list_date": job.dateofpublish,
            }
            jobs_list.append(job_info)

        # Get expired jobs
        # expiredyet = company_joblist.objects.filter(company=cid, Lastdate__gt=date.today())
        expiredyet = company_joblist.objects.filter(company=cid)
        # Convert string to list if stored incorrectly
        
        context = {
            "jobs": expiredyet,
            "jobs_list": jobs_list,
            "cname": cname,
            "compdt":company,
           
        }
        return render(request, 'company_jobs.html', context)
    
    except Exception as e:
        # Handle unexpected errors
        return render(request, 'company_jobs.html', {"cname": cname, 'compdt':company, "error": str(e)})
def jobseeker_notifications(request):
    jobseeker_id = request.session.get('jobseeker_id')

    if not jobseeker_id:
        return JsonResponse({"error": "User not logged in"}, status=400)

    notifications = JobNotification.objects.filter(jobseeker_profile__user=jobseeker_id, is_read=False).order_by('-created_at')

    notifications_data = [
        {
            "id": n.id,
            "message": n.message,
            "job_title": n.company_job.job_title.job_title,
            "company_name": n.company_job.company.name,
            "created_at": n.created_at.strftime("%d-%m-%Y"),
        }
        for n in notifications
    ]
    
    return JsonResponse({"notifications": notifications_data})

def mark_notifications_as_read(request):
    jobseeker_id = request.session.get('jobseeker_id')

    if not jobseeker_id:
        return JsonResponse({"error": "User not logged in"}, status=400)

    JobNotification.objects.filter(jobseeker_id=jobseeker_id, is_read=False).update(is_read=True)

    return JsonResponse({"message": "Notifications marked as read."})

def company_postjob(request):
    job = job_title.objects.all()
    cid = request.session.get('company_id')
    company = Company.objects.get(id=cid)
    if request.method == "POST":
        job_position = request.POST.get('job_title')
        job_description = request.POST.get('job_description')
        job_location = request.POST.get('Location')
        job_type = request.POST.get('Job_type')
        jobposted_date = request.POST.get('jobposted_date')
        qualification = request.POST.getlist('highest_qualification')
        percent_criteria=request.POST.get('percentage')
        skills = request.POST.getlist('skills_required')
        qualification_str = ",".join(qualification)
        skills_str = ",".join(skills)
        lastdate = request.POST.get('deadline')
        companyjob_obj = company_joblist()
        lastdate_obj = datetime.strptime(lastdate, "%Y-%m-%d").date()
        if lastdate_obj < now().date():
            messages.error(request,"Last date should not be past date.")
            redirect('job_listing')
        # ✅ **Check if this job_number already exists for this company**
        if company_joblist.objects.filter(company=cid, job_title=job_position).exists():
            messages.error(request, f"A job for job id `{job_position}` already exists in your listings. Please use a different Job ID.")
            return redirect('job_listing')
        companyjob_obj.company_id = cid
        companyjob_obj.job_title_id = job_position
        companyjob_obj.job_description = job_description
        companyjob_obj.location = job_location
        companyjob_obj.job_type = job_type
        companyjob_obj.dateofpublish = jobposted_date
        companyjob_obj.highest_qualification = qualification_str
        companyjob_obj.percent_criteria=percent_criteria
        companyjob_obj.skills_required = skills_str
        companyjob_obj.Lastdate = lastdate
        companyjob_obj.save()
        # **Notify jobseekers**
        jobseekers = jobseeker_profile.objects.all()
        for jobseeker in jobseekers:
            JobNotification.objects.create(
                    jobseeker_profile=jobseeker,
                    company_job=companyjob_obj,
                    message=f"New job posted: {companyjob_obj.job_title.job_title} at {company.name}."
                )
            messages.success(request, "Job Posted Successfully!!")        
    cname = request.session.get('company_name')
    return render(request, 'company_postjob.html', {"cname":cname, "compdt":company,'cdate':date.today().strftime("%Y-%m-%d"), "job":job})


def edit_job(request):
    job = job_title.objects.all()
    cid = request.session.get('company_id')
    cname = request.session.get('company_name')
    company = Company.objects.get(id=cid)
    if request.method == "POST":
        job_id = request.POST.get('job_id')
        job_description = request.POST.get('job_description')
        job_location = request.POST.get('location')
        job_type = request.POST.get('job_type')
        qualification = request.POST.getlist('highest_qualification')
        percent_criteria=request.POST.get('percentage')
        skills = request.POST.getlist('skills_required')
        qualification_str = ",".join(qualification)
        skills_str = ",".join(skills)
        lastdate = request.POST.get('deadline')
        companyjob_obj = company_joblist.objects.get(id=job_id)
        # **Retain old values if fields are not changed**
        job_type = job_type if job_type else companyjob_obj.job_type
        percent_criteria = percent_criteria if percent_criteria else companyjob_obj.percent_criteria
        # **Convert lastdate to date object and validate**
        if lastdate:
            try:
                lastdate_obj = datetime.strptime(lastdate, "%Y-%m-%d").date()
                if lastdate_obj < now().date():
                    messages.error(request, "The deadline cannot be set to a past date.")
                    return redirect(f'/edit_job/?ids={job_id}')  # Redirect with job ID
            except ValueError:
                messages.error(request, "Invalid date format. Please use YYYY-MM-DD.")
                return redirect(f'/edit_job/?ids={job_id}')
        else:
            lastdate_obj = companyjob_obj.Lastdate  # Retain old last date if not changed

     
        companyjob_obj.company_id = cid
        # companyjob_obj.job_number = job_number
        companyjob_obj.job_description = job_description
        companyjob_obj.location = job_location
        companyjob_obj.job_type = job_type
        companyjob_obj.dateofpublish = date.today().strftime('%Y-%m-%d')
        companyjob_obj.highest_qualification = qualification_str
        companyjob_obj.percent_criteria=percent_criteria
        companyjob_obj.skills_required = skills_str
        companyjob_obj.Lastdate = lastdate
        companyjob_obj.save()
        messages.success(request, "Job Edited Successfully!!") 
        return redirect('job_listing')
    jobid = request.GET.get('ids')
    jobobg = company_joblist.objects.get(id=jobid)
    context = {
        'jobobg':jobobg,
        "cname":cname,
        "job":job,
        "compdt":company,
       
    }
    return render(request, 'edit-job.html', context)

def applications(request):
    cid = request.session.get('company_id')
    cname = request.session.get('company_name')
    company = Company.objects.get(id=cid)
    Applicants=JobApplication.objects.filter(company_joblist__company_id=cid)
    # Sort the list of applicants by job title
    Applicants = sorted(Applicants, key=lambda x: x.company_joblist.job_title.job_title)
    # Group applicants by job title
    grouped_applicants = {}
    for job_title, group in groupby(Applicants, key=lambda x: x.company_joblist.job_title.job_title):
            grouped_applicants[job_title] = list(group)
    

    
    context = {
        "cname":cname,
        "compdt":company,
        "Applicants":Applicants,
        "grouped_applicants":grouped_applicants,
       
    }
    return render(request,'applications.html',context)

def candidates(request):
    cid = request.session.get('company_id')
    cname = request.session.get('company_name')
    company = Company.objects.get(id=cid)
    Applicants=JobApplication.objects.filter(company_joblist__company_id=cid,status="Accepted")
    # Sort the list of applicants by job title
    Applicants = sorted(Applicants, key=lambda x: x.company_joblist.job_title.job_title)
    # Group applicants by job title
    grouped_applicants = {}
    for job_title, group in groupby(Applicants, key=lambda x: x.company_joblist.job_title.job_title):
            grouped_applicants[job_title] = list(group)
    
    context = {
        "cname":cname,
        "compdt":company,
        "Applicants":Applicants,
        "grouped_applicants":grouped_applicants,
       
    }
    return render(request,'candidates.html',context)
def jobseeker_dashboard(request):
    fname = request.session.get('jobseeker_name')
    if not jobseeker_profile.objects.filter(name=fname).exists():
        return redirect('home')
    ccount = company_joblist.objects.all()
    profile=jobseeker_profile.objects.get(name=fname)
    profile_id=profile.id
    applied_jcount=JobApplication.objects.filter(jobseeker_id=profile_id)
    ajcount=applied_jcount.count()
    company = Company.objects.all()
    jcount = ccount.count()
    compcount = company.count()
    rejected=JobApplication.objects.filter(jobseeker_id__name=fname,status="Rejected").count()
    shortlisted=JobApplication.objects.filter(jobseeker_id__name=fname,status="Accepted").count()
    jobseeker = jobseeker_profile.objects.get(name=fname)
    unread_notifications = JobNotification.objects.filter(jobseeker_profile=jobseeker, is_read=False).count()
    return render(request, 'jobseeker_dashboard.html', {"fname":fname, "jcount":jcount, "compcount":compcount,"shortlisted_count":shortlisted,"rejected":rejected, "company":company, "jsdt":jobseeker,"ajcount":ajcount,"count":unread_notifications})

